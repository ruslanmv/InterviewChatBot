from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

def save_interview_history_old(history, language):
    """Saves the interview history to a TXT file."""
    file_name = f"interview_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    file_path = Path("hr_interviewer") / file_name

    with open(file_path, "w", encoding="utf-8") as file:
        for item in history:
            file.write(f"{item}\n")

    return file_path


import os
from datetime import datetime

def save_interview_history_fix(interview_history, language, folder_path="hr_interviewer"):
    """
    Saves the interview history to a file in the specified folder.

    Args:
        interview_history: The content of the interview history as a string.
        language: Language of the report.
        folder_path: Folder path where the history file will be saved.

    Returns:
        The file path of the saved interview history.
    """
    # Ensure the directory exists
    os.makedirs(folder_path, exist_ok=True)

    # Generate the filename with the current date and time
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = os.path.join(folder_path, f"interview_history_{timestamp}.txt")

    try:
        with open(file_path, "w", encoding="utf-8") as file:
            file.write("\n".join(interview_history))
        print(f"[DEBUG] Interview history saved at {file_path}")
        return file_path
    except Exception as e:
        print(f"[ERROR] Failed to save interview history: {e}")
        return None

import os
from datetime import datetime

def save_interview_history(interview_history, language, folder_path="hr_interviewer"):
    os.makedirs(folder_path, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = os.path.join(folder_path, f"interview_history_{timestamp}.txt")

    try:
        with open(file_path, "w", encoding="utf-8") as file:
            file.write("\n".join(interview_history))
        print(f"[DEBUG] Interview history saved at {file_path}")
        return file_path
    except Exception as e:
        print(f"[ERROR] Failed to save interview history: {e}")
        return None



def generate_interview_report(interview_history, language):
    """
    Generates a report in DOCX format based on the interview history.
    Args:
        interview_history: A list of strings representing the interview history.
        language: The language of the report.
    Returns:
        A tuple containing the report content as a string and the path to the generated DOCX file.
    """
    doc = Document()

    # Add title
    title = doc.add_heading('Interview Report', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(16)
    title_run.bold = True

    # Add date
    date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date_run = date_para.runs[0]
    date_run.font.name = 'Arial'
    date_run.font.size = Pt(11)

    # Add interview history
    doc.add_heading('Interview History', level=2)
    for item in interview_history:
        para = doc.add_paragraph(item)
        para_run = para.runs[0]
        para_run.font.name = 'Arial'
        para_run.font.size = Pt(12)

    # Save the document
    file_name = f"interview_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = Path("hr_interviewer") / file_name
    doc.save(file_path)

    # Convert DOCX to string (for display in Gradio, etc.)
    report_content = ""
    for para in doc.paragraphs:
        report_content += para.text + "\n"

    return report_content, file_path


import json
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

# ... (other functions remain the same)

def load_config(config_path="hr_interviewer/config.json"):
    """Loads the configuration from a JSON file."""
    try:
        with open(config_path, "r") as f:
            config = json.load(f)
    except FileNotFoundError:
        print(f"[WARNING] Config file not found at {config_path}. Using default settings.")
        config = {}  # Return empty dict to use defaults
    except json.JSONDecodeError:
        print(f"[ERROR] Error decoding JSON in {config_path}. Using default settings.")
        config = {}
    return config

def save_config(config, config_path="hr_interviewer/config.json"):
    """Saves the configuration to a JSON file."""
    try:
        with open(config_path, "w") as f:
            json.dump(config, f, indent=4)
        print(f"[INFO] Configuration saved to {config_path}")
        return True
    except Exception as e:
        print(f"[ERROR] Failed to save configuration: {e}")
        return False